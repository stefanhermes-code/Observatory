"""Read pricing documents to extract pricing structure."""
import sys
import os
sys.stdout.reconfigure(encoding='utf-8')

try:
    from docx import Document
    
    # Try reading the pricing document
    doc_path = os.path.join("Background Documentation", "Polyurethane Industry Observatory – Pricing.docx")
    print(f"Reading: {doc_path}")
    doc = Document(doc_path)
    
    print("=" * 80)
    print("PRICING DOCUMENT CONTENT:")
    print("=" * 80)
    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text)
    
    print("\n" + "=" * 80)
    print("TABLES:")
    print("=" * 80)
    for idx, table in enumerate(doc.tables):
        print(f"\n--- Table {idx + 1} ---")
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            if any(row_text):
                print(" | ".join(row_text))
    
except Exception as e:
    import traceback
    print(f"Error reading pricing doc: {e}")
    traceback.print_exc()

